import os
import base64
import requests
import json
import logging
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches

# ==========================================
# ENTERPRISE LEVEL OCR & VISION LLM PIPELINE
# PHASE 1.5/2: Tesseract + Ollama Integration
# ==========================================

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- CONFIGURATION ---
# IMPORTANT: Aapko Tesseract-OCR install karna hoga aur uska path yaha set karna hoga
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

OLLAMA_URL = "http://localhost:11434/api/generate"
VISION_MODEL = "llava"  # Make sure you run `ollama pull llava` in your terminal first

class VisionExtractor:
    def __init__(self, image_dir="extracted_images", output_docx="final_extracted_data.docx"):
        self.image_dir = image_dir
        self.output_docx = output_docx
        self.document = Document()
        self.document.add_heading("Enterprise Data Extraction (OCR + Vision LLM)", 0)

    def image_to_base64(self, image_path):
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def run_tesseract(self, image_path):
        """Runs Google's Tesseract OCR for high accuracy raw text extraction"""
        logger.info(f"Running Tesseract OCR on {os.path.basename(image_path)}")
        try:
            img = Image.open(image_path)
            text = pytesseract.image_to_string(img)
            return text.strip()
        except Exception as e:
            logger.error(f"Tesseract OCR Failed (Is it installed?): {e}")
            return f"[Tesseract Error: {e}]"

    def run_vision_llm(self, image_path):
        """Runs Ollama Vision LLM (LLaVA) for understanding layout and tables"""
        logger.info(f"Running Ollama ({VISION_MODEL}) on {os.path.basename(image_path)}")
        base64_img = self.image_to_base64(image_path)
        
        prompt = (
            "Analyze this image carefully. "
            "1. If you see a table, extract all data and format it precisely as a Markdown Table. "
            "2. If you see plain text, extract it accurately. "
            "3. If you see stamps or signatures, just mention '[Signature/Stamp Detected]'. "
            "Return only the extracted data without conversational filler."
        )

        payload = {
            "model": VISION_MODEL,
            "prompt": prompt,
            "images": [base64_img],
            "stream": False
        }

        try:
            response = requests.post(OLLAMA_URL, json=payload, timeout=300)
            if response.status_code == 200:
                result = response.json()
                return result.get("response", "").strip()
            else:
                logger.error(f"Ollama API Error: {response.text}")
                return "[Ollama LLM Error: Failed to process image]"
        except requests.exceptions.ConnectionError:
            logger.error("Could not connect to Ollama. Make sure Ollama app is running.")
            return "[Connection Error: Ollama is not running]"
        except Exception as e:
            logger.error(f"Ollama Exception: {e}")
            return f"[Ollama Exception: {e}]"

    def process_all_images(self):
        if not os.path.exists(self.image_dir):
            logger.error(f"Directory {self.image_dir} not found!")
            return

        # Sort images to process them in correct page order
        images = sorted(
            [f for f in os.listdir(self.image_dir) if f.endswith(('.png', '.jpeg', '.jpg'))],
            key=lambda x: int(x.split('_')[1]) if '_' in x else 0
        )

        if not images:
            logger.warning(f"No images found in {self.image_dir}")
            return

        for img_name in images:
            img_path = os.path.join(self.image_dir, img_name)
            logger.info(f"--- Processing {img_name} ---")
            
            self.document.add_heading(f"Source: {img_name}", level=1)
            
            # Insert original image for reference
            try:
                self.document.add_picture(img_path, width=Inches(5.0))
            except Exception:
                pass

            # 1. Tesseract Output
            self.document.add_heading("Tesseract OCR Output (Raw Text)", level=2)
            tess_text = self.run_tesseract(img_path)
            self.document.add_paragraph(tess_text if tess_text else "[No text detected by Tesseract]")

            # 2. Vision LLM Output
            self.document.add_heading("Ollama Vision LLM Output (Tables & Structure)", level=2)
            llm_text = self.run_vision_llm(img_path)
            self.document.add_paragraph(llm_text if llm_text else "[No structure detected by LLM]")
            
            self.document.add_page_break()

        # Save Final Document
        self.document.save(self.output_docx)
        logger.info(f"Successfully processed all images. Final report saved to: {self.output_docx}")

if __name__ == "__main__":
    extractor = VisionExtractor(
        image_dir=r"C:\Users\anshu\Desktop\extract\extracted_images",
        output_docx=r"C:\Users\anshu\Desktop\extract\final_extracted_data.docx"
    )
    extractor.process_all_images()
