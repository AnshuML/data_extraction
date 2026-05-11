import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Inches
import os
import logging
from PIL import Image
import io

# ==========================================
# ENTERPRISE LEVEL PDF EXTRACTION PIPELINE
# PHASE 1: Open Source Tooling Integration
# ==========================================

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class OCREngine:
    """
    Placeholder for Open Source OCR / LLM Engine (e.g., Tesseract, LLaVA, or Qwen-VL via Ollama)
    Since the current PDF is scanned (images only), this engine will be needed for Phase 1.5.
    """
    def __init__(self):
        # Setup Tesseract or Ollama client here
        pass
        
    def process_image(self, image_path):
        # Implementation for passing image to Open Source LLM / OCR
        # return extracted_text, extracted_tables
        return "OCR/LLM text extraction pending...", []


class PDFExtractor:
    def __init__(self, pdf_path, output_docx_path, image_output_dir="extracted_images"):
        self.pdf_path = pdf_path
        self.output_docx_path = output_docx_path
        self.image_output_dir = image_output_dir
        self.ocr_engine = OCREngine()
        
        if not os.path.exists(self.image_output_dir):
            os.makedirs(self.image_output_dir)
            
        self.document = Document()
        self.document.add_heading("Automated PDF Data Extraction Report", 0)
        
    def extract_all(self):
        logger.info(f"Starting enterprise extraction pipeline for: {self.pdf_path}")
        
        try:
            pdf_fitz = fitz.open(self.pdf_path)
            pdf_plumber = pdfplumber.open(self.pdf_path)
            
            num_pages = len(pdf_fitz)
            logger.info(f"Total pages detected: {num_pages}")
            
            for page_num in range(num_pages):
                logger.info(f"Processing Page {page_num + 1}/{num_pages}")
                self.document.add_heading(f"Page {page_num + 1} Data", level=1)
                
                # 1. Native Digital Text Extraction
                text_found = self._extract_text(pdf_fitz[page_num])
                
                # 2. Native Digital Table Extraction
                tables_found = self._extract_tables(pdf_plumber.pages[page_num])
                
                # 3. Image & Stamp Extraction (Critical for scanned PDFs)
                images_extracted = self._extract_images(pdf_fitz, page_num)
                
                # 4. Fallback / LLM Processing for Scanned Pages
                if not text_found and not tables_found and images_extracted:
                    logger.info("Page appears to be scanned. Initiating Open Source LLM / OCR fallback pipeline (Pending Setup).")
                    self.document.add_paragraph("Note: Page is scanned. Passing extracted images to Open Source LLM/OCR pipeline for text/table detection.")
                
                self.document.add_page_break()
                
            self.document.save(self.output_docx_path)
            logger.info(f"Extraction Pipeline Complete. Saved to: {self.output_docx_path}")
            
        except Exception as e:
            logger.error(f"Pipeline Failure: {e}")
        finally:
            if 'pdf_fitz' in locals(): pdf_fitz.close()
            if 'pdf_plumber' in locals(): pdf_plumber.close()

    def _extract_text(self, fitz_page):
        text = fitz_page.get_text("text").strip()
        if text:
            self.document.add_heading("Extracted Text", level=2)
            self.document.add_paragraph(text)
            return True
        return False

    def _extract_tables(self, plumber_page):
        tables = plumber_page.extract_tables()
        if tables:
            for table_idx, table in enumerate(tables):
                self.document.add_heading(f"Extracted Table {table_idx + 1}", level=2)
                if not table or not table[0]: continue
                    
                word_table = self.document.add_table(rows=len(table), cols=len(table[0]))
                word_table.style = 'Table Grid'
                
                for row_idx, row in enumerate(table):
                    for col_idx, cell in enumerate(row):
                        if cell is not None:
                            try:
                                word_table.cell(row_idx, col_idx).text = str(cell).replace('\n', ' ')
                            except IndexError:
                                pass
            self.document.add_paragraph("\n")
            return True
        return False

    def _extract_images(self, pdf_fitz, page_num):
        image_list = pdf_fitz[page_num].get_images(full=True)
        images_saved = []
        
        if image_list:
            self.document.add_heading("Extracted Images / Stamps / Scanned Pages", level=2)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_fitz.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                image_name = f"page_{page_num+1}_img_{img_index+1}.{image_ext}"
                image_path = os.path.join(self.image_output_dir, image_name)
                
                with open(image_path, "wb") as f:
                    f.write(image_bytes)
                images_saved.append(image_path)
                
                try:
                    self.document.add_picture(image_path, width=Inches(6.0))
                except Exception as e:
                    self.document.add_paragraph(f"[Image saved to: {image_path}]")
                    
        return len(images_saved) > 0

if __name__ == "__main__":
    pdf_file_path = r"C:\Users\anshu\Desktop\extract\data\Balance Sheet of DSL 114045 (1).pdf"
    word_file_path = r"C:\Users\anshu\Desktop\extract\output_extracted_data.docx"
    
    extractor = PDFExtractor(pdf_file_path, word_file_path)
    extractor.extract_all()
