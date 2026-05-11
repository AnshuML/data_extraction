import os
import glob
import fitz
import base64
import requests
import pytesseract
import logging
import shutil
from PIL import Image
from docx import Document
from docx.shared import Inches

# ==========================================
# ENTERPRISE BATCH PROCESSING WRAPPER
# Fully Automated Pipeline for Multiple Files
# ==========================================

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- ENTERPRISE CONFIGURATION ---
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
OLLAMA_URL = "http://localhost:11434/api/generate"
VISION_MODEL = "llava"

class BatchPDFProcessor:
    def __init__(self, data_dir="data", output_dir="outputs"):
        self.data_dir = data_dir
        self.output_dir = output_dir
        
        # Ensure output directory exists
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def image_to_base64(self, image_path):
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def extract_images_from_pdf(self, pdf_path, image_output_dir):
        if not os.path.exists(image_output_dir):
            os.makedirs(image_output_dir)
            
        pdf_fitz = fitz.open(pdf_path)
        images_extracted = []
        
        for page_num in range(len(pdf_fitz)):
            image_list = pdf_fitz[page_num].get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_fitz.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                image_name = f"page_{page_num+1}_img_{img_index+1}.{image_ext}"
                image_path = os.path.join(image_output_dir, image_name)
                
                with open(image_path, "wb") as f:
                    f.write(image_bytes)
                images_extracted.append(image_path)
                
        pdf_fitz.close()
        return images_extracted

    def run_tesseract(self, image_path):
        try:
            img = Image.open(image_path)
            return pytesseract.image_to_string(img).strip()
        except Exception as e:
            return f"[Tesseract Error: {e}]"

    def run_vision_llm(self, image_path):
        base64_img = self.image_to_base64(image_path)
        prompt = (
            "Analyze this image carefully. "
            "1. If you see a table, extract all data and format it precisely as a Markdown Table. "
            "2. If you see plain text, extract it accurately. "
            "3. If you see stamps or signatures, just mention '[Signature/Stamp Detected]'. "
            "Return only the extracted data without conversational filler."
        )

        payload = {
            "model": VISION_MODEL,
            "prompt": prompt,
            "images": [base64_img],
            "stream": False
        }

        try:
            response = requests.post(OLLAMA_URL, json=payload, timeout=300)
            if response.status_code == 200:
                return response.json().get("response", "").strip()
            else:
                return "[Ollama API Error]"
        except Exception as e:
            return f"[Ollama Connection Error: {e}]"

    def process_single_file(self, pdf_path):
        filename = os.path.basename(pdf_path)
        base_name = os.path.splitext(filename)[0]
        
        logger.info(f"========== STARTING PROCESSING: {filename} ==========")
        
        # Dedicated folders and files for this specific PDF
        doc_output_path = os.path.join(self.output_dir, f"Report_{base_name}.docx")
        temp_img_dir = os.path.join(self.output_dir, f"temp_images_{base_name}")
        
        document = Document()
        document.add_heading(f"Data Extraction Report: {filename}", 0)
        
        try:
            # 1. Extract Images
            logger.info(f"Extracting images from {filename}...")
            extracted_images = self.extract_images_from_pdf(pdf_path, temp_img_dir)
            
            if not extracted_images:
                logger.warning(f"No images found in {filename}. Skipping Vision pipeline.")
                document.add_paragraph("No scannable images found in this PDF.")
            
            # 2. Process Each Image (OCR + LLM)
            for img_path in extracted_images:
                img_name = os.path.basename(img_path)
                logger.info(f"Processing {img_name}...")
                document.add_heading(f"Source: {img_name}", level=1)
                
                try:
                    document.add_picture(img_path, width=Inches(5.0))
                except:
                    pass
                
                document.add_heading("Tesseract OCR Output", level=2)
                tess_text = self.run_tesseract(img_path)
                document.add_paragraph(tess_text if tess_text else "[No text]")

                document.add_heading("Ollama Vision LLM Output", level=2)
                llm_text = self.run_vision_llm(img_path)
                document.add_paragraph(llm_text if llm_text else "[No data]")
                
                document.add_page_break()
                
            # 3. Save Final Document for this PDF
            document.save(doc_output_path)
            logger.info(f"SUCCESS: Report saved to {doc_output_path}")
            
        except Exception as e:
            logger.error(f"FAILED to process {filename}: {e}")
        finally:
            # Enterprise Cleanup: Remove temporary images to save disk space
            if os.path.exists(temp_img_dir):
                shutil.rmtree(temp_img_dir)
                logger.info(f"Cleaned up temporary images for {filename}")

    def run_batch(self):
        # Dynamically find ALL pdf files in the data directory
        pdf_files = glob.glob(os.path.join(self.data_dir, "*.pdf"))
        if not pdf_files:
            logger.error(f"No PDF files found in '{self.data_dir}' folder.")
            return
            
        logger.info(f"Found {len(pdf_files)} PDF files to process in batch.")
        
        # Process them one by one automatically
        for idx, pdf_path in enumerate(pdf_files):
            logger.info(f"\n--- Processing File {idx+1} of {len(pdf_files)} ---")
            self.process_single_file(pdf_path)
            
        logger.info("========== BATCH PROCESSING COMPLETE ==========")

if __name__ == "__main__":
    # Point data_dir to where the PDFs live. Output will go to 'outputs' folder.
    processor = BatchPDFProcessor(
        data_dir=r"C:\Users\anshu\Desktop\extract\data",
        output_dir=r"C:\Users\anshu\Desktop\extract\outputs"
    )
    processor.run_batch()
